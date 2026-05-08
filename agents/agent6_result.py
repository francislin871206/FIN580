"""
==============================================================================
  AGENT 6 -- RESULT ANALYSIS AGENT (Report Generator)
  FIN580 Quantamental Investment Project -- Brent Crude Oil
==============================================================================
  Reads all outputs from Agents 1-5 and uses Gemini to synthesize
  a comprehensive, professional-grade analysis report in Markdown format.
  Usage: python -m agents.agent6_result
  Output: data/processed/analysis_report.md, data/processed/analysis_report.docx
==============================================================================
"""

import os, json, time, re
from agents import load_config, load_prompt, resolve_path

_CFG = load_config()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL = _CFG.get("gemini", {}).get("model", "gemini-2.5-flash")

def load_agent_outputs():
    """Load all 5 agent output files from data/processed/."""
    paths = _CFG.get("paths", {})
    keys = {"a1_events":"a1_events","a2_macro":"a2_macro","a3_signals":"a3_signals","a4_risk":"a4_risk","a5_trades":"a5_trades"}
    data = {}
    for key, cfg_key in keys.items():
        p = str(resolve_path(paths.get(cfg_key, f"data/processed/{key}.json")))
        if os.path.exists(p):
            with open(p,"r",encoding="utf-8") as f: data[key] = json.load(f)
            print(f"  [OK] Loaded {p}")
        else: print(f"  [SKIP] {p} not found"); data[key] = []
    return data

def build_summaries(data):
    a1 = [{"event_id":e.get("event_id"),"headline":e.get("headline"),"event_type":e.get("event_type"),"directional_bias":e.get("directional_bias"),"confidence":e.get("confidence_score"),"sentiment":e.get("sentiment_score"),"source":e.get("source"),"timestamp":e.get("timestamp")} for e in data.get("a1_events",[])]
    a2 = [{"event_id":m.get("event_id"),"macro_thesis":m.get("macro_thesis"),"causal_chain":m.get("causal_chain"),"market_regime":m.get("market_regime"),"expected_impact":m.get("expected_impact"),"conviction_score":m.get("conviction_score"),"time_horizon":m.get("time_horizon")} for m in data.get("a2_macro",[])]
    a3 = data.get("a3_signals",[])
    a4 = [{"event_id":d.get("signal",{}).get("event_id"),"direction":d.get("signal",{}).get("direction"),"conviction":d.get("signal",{}).get("conviction"),"approved":d.get("approved"),"veto_reason":d.get("veto_reason"),"risk_multiplier":d.get("risk_multiplier")} for d in data.get("a4_risk",[])]
    a5 = data.get("a5_trades",[])
    return a1, a2, a3, a4, a5

def generate_report_with_gemini(a1, a2, a3, a4, a5):
    prompt = load_prompt("report_generation.txt",
        A1_DATA=json.dumps(a1,indent=2), A2_DATA=json.dumps(a2,indent=2),
        A3_DATA=json.dumps(a3,indent=2), A4_DATA=json.dumps(a4,indent=2),
        A5_DATA=json.dumps(a5,indent=2))
    from google import genai; from google.genai import types
    client = genai.Client(api_key=GEMINI_API_KEY)
    print("  [INFO] Sending data to Gemini for report generation...")
    resp = client.models.generate_content(model=GEMINI_MODEL, contents=prompt,
        config=types.GenerateContentConfig(system_instruction="You are an expert quantitative investment analyst. Write a professional Markdown report with tables and specific numbers."))
    return resp.text

def generate_fallback_report(a1, a2, a3, a4, a5):
    lines = [
        "# Multi-Agent Trading System -- Analysis Report",
        f"**Generated**: {time.strftime('%Y-%m-%d %H:%M:%S')}","","---","",
        "## 1. Executive Summary",
        f"- **Events detected**: {len(a1)}",f"- **Macro analyses**: {len(a2)}",
        f"- **Signals generated**: {len(a3)}",f"- **Risk decisions**: {len(a4)}",
        f"- **Trades executed**: {len(a5)}","","---","",
        "## 2. Agent 1: Event Detection","",
        "| # | Headline | Type | Bias | Confidence | Sentiment | Source |",
        "|---|----------|------|------|------------|-----------|--------|"]
    for i, ev in enumerate(a1):
        lines.append(f"| {i+1} | {ev['headline'][:60]}... | {ev['event_type']} | {ev['directional_bias']} | {ev['confidence']:.2f} | {ev['sentiment']:.2f} | {ev['source']} |")
    lines += ["","---","","## 3. Agent 2: Macro Debate Conclusions",""]
    for m in a2:
        lines += [f"### Event: {m['event_id']}",f"- **Thesis**: {m['macro_thesis']}",f"- **Regime**: {m['market_regime']}",f"- **Impact**: {m['expected_impact']}",f"- **Conviction**: {m['conviction_score']}/5",""]
    lines += ["---","","## 4. Agent 3: Trading Signals",""]
    lines.append(f"- LONG: {sum(1 for s in a3 if s.get('direction')=='LONG')} | SHORT: {sum(1 for s in a3 if s.get('direction')=='SHORT')} | FLAT: {sum(1 for s in a3 if s.get('direction')=='FLAT')}")
    lines += ["","---","","## 5. Agent 4: Risk Gating",""]
    approved = sum(1 for d in a4 if d.get("approved"))
    lines += [f"- **Approved**: {approved} | **Vetoed**: {len(a4)-approved}","",
        "| Event ID | Direction | Conviction | Status | Reason |",
        "|----------|-----------|------------|--------|--------|"]
    for d in a4: lines.append(f"| {d['event_id']} | {d['direction']} | {d['conviction']} | {'APPROVED' if d['approved'] else 'VETOED'} | {d['veto_reason']} |")
    lines += ["","---","","## 6. Agent 5: Portfolio Trades",""]
    total = sum(t.get("size_usd",0) for t in a5)
    lines.append(f"**Total Capital Deployed**: ${total:,.2f}")
    for t in a5: lines += [f"- **{t['date']}**: {t['direction']} {t['asset']} -- ${t['size_usd']:,.2f} @ ${t['entry_price_ref']:.2f}",f"  - {t['reasoning_trace']}"]
    lines += ["","---","","## 7. Risk Assessment","",
        "- Model relies on mock VIX data for risk gating in standalone mode.",
        "- Historical analog comparisons use static benchmarks.",
        "- Cash buffer constraints may over-restrict position sizing.",
        "","---","","## 8. Recommendations","",
        "- Integrate live VIX and drawdown data for production.",
        "- Expand news sources beyond NewsAPI.",
        "- Run ablation studies to measure each agent's contribution."]
    return "\n".join(lines)

def save_report_to_docx(md_text, output_path):
    """Convert Markdown report to a Word document."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)
    for lv in range(1,5): doc.styles[f"Heading {lv}"].font.color.rgb = RGBColor(0x1A,0x3C,0x6E)

    def add_rich(p, text):
        for part in re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text):
            if not part: continue
            if part.startswith("**") and part.endswith("**"): r=p.add_run(part[2:-2]); r.bold=True
            elif part.startswith("*") and part.endswith("*"): r=p.add_run(part[1:-1]); r.italic=True
            elif part.startswith("`") and part.endswith("`"): r=p.add_run(part[1:-1]); r.font.name="Consolas"; r.font.size=Pt(10)
            else: p.add_run(part)

    def flush_table(tlines):
        if len(tlines)<2: return
        hdr = [c.strip() for c in tlines[0].strip("|").split("|")]
        rows = [[c.strip() for c in r.strip("|").split("|")] for r in tlines[2:]]
        nc = len(hdr); t = doc.add_table(rows=1+len(rows),cols=nc)
        t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(hdr): t.rows[0].cells[i].text=""; r=t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
        for ri,rd in enumerate(rows):
            for ci,cd in enumerate(rd):
                if ci<nc: t.rows[ri+1].cells[ci].text=""; t.rows[ri+1].cells[ci].paragraphs[0].add_run(cd).font.size=Pt(9)
        doc.add_paragraph()

    lines = md_text.replace("\r\n","\n").split("\n"); tbuf=[]; i=0
    while i<len(lines):
        ln = lines[i]
        if ln.strip().startswith("|"): tbuf.append(ln); i+=1; continue
        else:
            if tbuf: flush_table(tbuf); tbuf=[]
        s = ln.strip()
        if not s: i+=1; continue
        if re.match(r'^-{3,}$',s) or re.match(r'^\*{3,}$',s):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run("─"*60); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC); r.font.size=Pt(8); i+=1; continue
        hm = re.match(r'^(#{1,4})\s+(.+)$',s)
        if hm: doc.add_heading(hm.group(2),level=len(hm.group(1))); i+=1; continue
        bm = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$',s)
        if bm:
            p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(0.25+0.25*min(len(bm.group(1))//4,2))
            add_rich(p,bm.group(3)); i+=1; continue
        p=doc.add_paragraph(); add_rich(p,s); i+=1
    if tbuf: flush_table(tbuf)
    doc.save(output_path); print(f"  [OK] Word document saved to: {output_path}")

def main():
    print("="*70+"\n  AGENT 6 -- RESULT ANALYSIS (Report Generator)\n"+"="*70)
    data = load_agent_outputs(); a1,a2,a3,a4,a5 = build_summaries(data)
    try:
        report = generate_report_with_gemini(a1,a2,a3,a4,a5)
        print(f"\n  [OK] Gemini report generated ({len(report):,} chars)")
    except Exception as e:
        print(f"\n  [WARN] Gemini failed: {e}\n  [INFO] Using fallback report...")
        report = generate_fallback_report(a1,a2,a3,a4,a5)
    rp = str(resolve_path(_CFG["paths"]["report_md"]))
    os.makedirs(os.path.dirname(rp),exist_ok=True)
    with open(rp,"w",encoding="utf-8") as f: f.write(report)
    print(f"  [OK] Markdown report saved to: {rp}")
    dp = str(resolve_path(_CFG["paths"]["report_docx"]))
    try: save_report_to_docx(report, dp)
    except ImportError: print("  [WARN] python-docx not installed.")
    except Exception as e: print(f"  [WARN] Word export failed: {e}")
    for ln in report.split("\n")[:15]: print(f"  {ln}")

if __name__ == "__main__": main()
