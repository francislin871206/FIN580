"""
==============================================================================
  AGENT 6 -- RESULT ANALYSIS AGENT (Report Generator)
  FIN580 Quantamental Investment Project -- Brent Crude Oil
==============================================================================
  This agent reads all outputs from Agents 1-5 and uses Gemini to synthesize
  a comprehensive, professional-grade analysis report in Markdown format.

  Pipeline Position:
    A1 -> A2 -> A3 -> A4 -> A5 -> **A6 (Result Agent)** -> analysis_report.md

  Usage:
    python result_agent.py

  Output:
    data/processed/analysis_report.md
==============================================================================
"""

import os
import json
import time
import re

# ── API KEY ─────────────────────────────────────────────────────────────────
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyChV0UY6NcbVDT-O5ukXFjPOLIfcTSrknE")
GEMINI_MODEL = "gemini-2.5-flash"


def load_agent_outputs():
    """Load all 5 agent output files from data/processed/."""
    agent_files = {
        "a1_events":  "data/processed/a1_events.json",
        "a2_macro":   "data/processed/a2_macro.json",
        "a3_signals": "data/processed/a3_signals.json",
        "a4_risk":    "data/processed/a4_risk_decisions.json",
        "a5_trades":  "data/processed/a5_portfolio_trades.json",
    }

    agent_data = {}
    for key, path in agent_files.items():
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                agent_data[key] = json.load(f)
            print(f"  [OK] Loaded {path}")
        else:
            print(f"  [SKIP] {path} not found")
            agent_data[key] = []

    return agent_data


def build_summaries(agent_data):
    """Build clean summaries for each agent (strip long debate transcripts)."""

    # --- A1: Event Summary ---
    a1 = []
    for ev in agent_data.get("a1_events", []):
        a1.append({
            "event_id":         ev.get("event_id"),
            "headline":         ev.get("headline"),
            "event_type":       ev.get("event_type"),
            "directional_bias": ev.get("directional_bias"),
            "confidence":       ev.get("confidence_score"),
            "sentiment":        ev.get("sentiment_score"),
            "source":           ev.get("source"),
            "timestamp":        ev.get("timestamp"),
        })

    # --- A2: Macro Summary (exclude full debate transcript) ---
    a2 = []
    for m in agent_data.get("a2_macro", []):
        a2.append({
            "event_id":        m.get("event_id"),
            "macro_thesis":    m.get("macro_thesis"),
            "causal_chain":    m.get("causal_chain"),
            "market_regime":   m.get("market_regime"),
            "expected_impact": m.get("expected_impact"),
            "conviction_score": m.get("conviction_score"),
            "time_horizon":    m.get("time_horizon"),
        })

    # --- A3: Signal Summary ---
    a3 = agent_data.get("a3_signals", [])

    # --- A4: Risk Summary ---
    a4 = []
    for d in agent_data.get("a4_risk", []):
        a4.append({
            "event_id":        d.get("signal", {}).get("event_id"),
            "direction":       d.get("signal", {}).get("direction"),
            "conviction":      d.get("signal", {}).get("conviction"),
            "approved":        d.get("approved"),
            "veto_reason":     d.get("veto_reason"),
            "risk_multiplier": d.get("risk_multiplier"),
        })

    # --- A5: Trades Summary ---
    a5 = agent_data.get("a5_trades", [])

    return a1, a2, a3, a4, a5


def generate_report_with_gemini(a1, a2, a3, a4, a5):
    """Call Gemini to write a professional analysis report."""

    prompt = f"""
You are a senior quantitative strategist writing a professional investment analysis report
for a multi-agent trading system that analyzes Brent Crude Oil markets.

The system has 5 agents that process data sequentially:
- Agent 1 (Event Detection): Detects and classifies news events using FinBERT + spaCy + Gemini
- Agent 2 (Macro Interpretation): Runs a 3-round adversarial debate to form macro theses
- Agent 3 (Signal Generation): Converts macro theses into LONG/SHORT/FLAT trading signals
- Agent 4 (Risk Management): Applies VIX, drawdown, and conviction checks to approve/veto trades
- Agent 5 (Portfolio Construction): Sizes approved trades with cash buffer constraints

Below is the structured output from each agent. Please write a comprehensive analysis report
in Markdown format that includes:

1. **Executive Summary** - Key findings and overall market view
2. **Agent 1: Event Detection Results** - What events were detected, sentiment breakdown, source diversity
3. **Agent 2: Macro Debate Conclusions** - Key macro theses, where the debating agents agreed/disagreed, conviction levels
4. **Agent 3: Trading Signals** - Signal direction breakdown (how many LONG/SHORT/FLAT), conviction distribution
5. **Agent 4: Risk Gating Results** - How many signals were approved vs vetoed, reasons for vetoes
6. **Agent 5: Final Portfolio Decisions** - Executed trades, total capital deployed, position sizing rationale
7. **Risk Assessment & Limitations** - Key risks to the current positions, model limitations, areas for improvement
8. **Recommendations** - Actionable next steps for the trading team

Use tables where appropriate. Be specific with numbers. Write in a professional, institutional tone.
The report should be suitable for presentation to a portfolio manager or academic review committee.

=== AGENT 1 OUTPUT (Event Detection) ===
{json.dumps(a1, indent=2)}

=== AGENT 2 OUTPUT (Macro Interpretation) ===
{json.dumps(a2, indent=2)}

=== AGENT 3 OUTPUT (Trading Signals) ===
{json.dumps(a3, indent=2)}

=== AGENT 4 OUTPUT (Risk Decisions) ===
{json.dumps(a4, indent=2)}

=== AGENT 5 OUTPUT (Portfolio Trades) ===
{json.dumps(a5, indent=2)}
"""

    from google import genai
    from google.genai import types

    client = genai.Client(api_key=GEMINI_API_KEY)
    print("  [INFO] Sending data to Gemini for report generation...")

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            system_instruction=(
                "You are an expert quantitative investment analyst at a top-tier hedge fund. "
                "Write a professional, detailed analysis report in Markdown. "
                "Use tables, headers, and bullet points for clarity. "
                "Be specific with numbers and percentages from the data provided. "
                "The report should demonstrate deep understanding of oil markets and risk management."
            ),
        ),
    )

    return response.text


def generate_fallback_report(a1, a2, a3, a4, a5):
    """Generate a basic structured report without Gemini (offline fallback)."""

    lines = [
        "# Multi-Agent Trading System -- Analysis Report",
        f"**Generated**: {time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "---",
        "",
        "## 1. Executive Summary",
        f"- **Events detected**: {len(a1)}",
        f"- **Macro analyses completed**: {len(a2)}",
        f"- **Trading signals generated**: {len(a3)}",
        f"- **Risk decisions made**: {len(a4)}",
        f"- **Trades executed**: {len(a5)}",
        "",
        "---",
        "",
        "## 2. Agent 1: Event Detection",
        "",
        "| # | Headline | Type | Bias | Confidence | Sentiment | Source |",
        "|---|----------|------|------|------------|-----------|--------|",
    ]
    for i, ev in enumerate(a1):
        lines.append(
            f"| {i+1} | {ev['headline'][:60]}... | {ev['event_type']} | "
            f"{ev['directional_bias']} | {ev['confidence']:.2f} | "
            f"{ev['sentiment']:.2f} | {ev['source']} |"
        )

    lines += ["", "---", "", "## 3. Agent 2: Macro Debate Conclusions", ""]
    for m in a2:
        lines.append(f"### Event: {m['event_id']}")
        lines.append(f"- **Thesis**: {m['macro_thesis']}")
        lines.append(f"- **Market Regime**: {m['market_regime']}")
        lines.append(f"- **Expected Impact**: {m['expected_impact']}")
        lines.append(f"- **Conviction**: {m['conviction_score']}/5")
        lines.append(f"- **Time Horizon**: {m['time_horizon']}")
        lines.append("")

    lines += ["---", "", "## 4. Agent 3: Trading Signals", ""]
    long_count  = sum(1 for s in a3 if s.get("direction") == "LONG")
    short_count = sum(1 for s in a3 if s.get("direction") == "SHORT")
    flat_count  = sum(1 for s in a3 if s.get("direction") == "FLAT")
    lines.append(f"- LONG: {long_count} | SHORT: {short_count} | FLAT: {flat_count}")
    lines.append("")

    lines += ["---", "", "## 5. Agent 4: Risk Gating", ""]
    approved = sum(1 for d in a4 if d.get("approved"))
    vetoed   = len(a4) - approved
    lines.append(f"- **Approved**: {approved} | **Vetoed**: {vetoed}")
    lines.append("")
    lines.append("| Event ID | Direction | Conviction | Status | Reason |")
    lines.append("|----------|-----------|------------|--------|--------|")
    for d in a4:
        status = "APPROVED" if d["approved"] else "VETOED"
        lines.append(f"| {d['event_id']} | {d['direction']} | {d['conviction']} | {status} | {d['veto_reason']} |")

    lines += ["", "---", "", "## 6. Agent 5: Portfolio Trades", ""]
    total = sum(t.get("size_usd", 0) for t in a5)
    lines.append(f"**Total Capital Deployed**: ${total:,.2f}")
    lines.append("")
    for t in a5:
        lines.append(f"- **{t['date']}**: {t['direction']} {t['asset']} -- ${t['size_usd']:,.2f} @ ${t['entry_price_ref']:.2f}")
        lines.append(f"  - Reasoning: {t['reasoning_trace']}")

    lines += ["", "---", "", "## 7. Risk Assessment", ""]
    lines.append("- Model relies on mock VIX data for risk gating in standalone mode.")
    lines.append("- Historical analog comparisons use static benchmarks.")
    lines.append("- Cash buffer constraints may over-restrict position sizing.")

    lines += ["", "---", "", "## 8. Recommendations", ""]
    lines.append("- Integrate live VIX and drawdown data for production risk gating.")
    lines.append("- Expand news sources beyond NewsAPI for broader event coverage.")
    lines.append("- Run ablation studies (remove A1, A2, or A4) to measure each agent's contribution.")

    return "\n".join(lines)


def save_report_to_docx(markdown_text, output_path):
    """
    Convert a Markdown report string into a professionally formatted Word
    document (.docx) using python-docx.

    Supports:
      - Heading levels 1-4 (# through ####)
      - Bold (**text**) and italic (*text*) inline formatting
      - Unordered bullet lists (-, *, numbered)
      - Markdown tables (| col1 | col2 | ... |)
      - Horizontal rules (---) rendered as a page break
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── Global style tweaks ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)   # dark navy

    # ── Helper: apply bold / italic runs to a paragraph ────────────────────
    def add_rich_text(paragraph, text):
        """Parse **bold**, *italic*, `code` and add as runs."""
        # Pattern: **bold**, *italic*, `code`, or plain text
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
            else:
                paragraph.add_run(part)

    # ── Helper: add a styled table from Markdown table lines ──────────────
    def flush_table(table_lines):
        """Parse markdown table lines and add a Word table."""
        if len(table_lines) < 2:
            return
        # Parse header
        header_cells = [c.strip() for c in table_lines[0].strip("|").split("|")]
        # Skip separator row (index 1), parse data rows
        data_rows = []
        for row_line in table_lines[2:]:
            cells = [c.strip() for c in row_line.strip("|").split("|")]
            data_rows.append(cells)

        num_cols = len(header_cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, text in enumerate(header_cells):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)

        # Data rows
        for r_idx, row_data in enumerate(data_rows):
            for c_idx, text in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    run.font.size = Pt(9)

        # Add a blank paragraph after the table for spacing
        doc.add_paragraph()

    # ── Main parsing loop ─────────────────────────────────────────────────
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    table_buffer = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Table detection: lines starting with | ────────────────────────
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        else:
            # Flush any buffered table
            if table_buffer:
                flush_table(table_buffer)
                table_buffer = []

        stripped = line.strip()

        # ── Blank line ────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule (---) ─────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            # Add a subtle divider paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            p = doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # ── Bullet / list items ───────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$', stripped)
        if bullet_match:
            indent = len(bullet_match.group(1))
            content = bullet_match.group(3)
            # Determine nesting level (0 or 1)
            nest_level = min(indent // 4, 2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * nest_level)
            add_rich_text(p, content)
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────
        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    # Flush any remaining table
    if table_buffer:
        flush_table(table_buffer)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"  [OK] Word document saved to: {output_path}")


def main():
    print("=" * 70)
    print("  AGENT 6 -- RESULT ANALYSIS (Gemini Report Generator)")
    print("=" * 70)

    # 1. Load all agent outputs
    agent_data = load_agent_outputs()
    a1, a2, a3, a4, a5 = build_summaries(agent_data)

    # 2. Generate report
    try:
        report_text = generate_report_with_gemini(a1, a2, a3, a4, a5)
        print(f"\n  [OK] Gemini report generated ({len(report_text):,} characters)")
    except Exception as e:
        print(f"\n  [WARN] Gemini API failed: {e}")
        print("  [INFO] Falling back to offline report generator...")
        report_text = generate_fallback_report(a1, a2, a3, a4, a5)

    # 3. Save Markdown report
    os.makedirs("data/processed", exist_ok=True)
    report_path = "data/processed/analysis_report.md"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_text)
    print(f"  [OK] Markdown report saved to: {report_path}")

    # 4. Save Word document
    docx_path = "data/processed/analysis_report.docx"
    try:
        save_report_to_docx(report_text, docx_path)
    except ImportError:
        print("  [WARN] python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        print(f"  [WARN] Word export failed: {e}")

    # 5. Print preview
    preview = report_text.split("\n")[:25]
    print("\n  --- REPORT PREVIEW ---")
    for line in preview:
        print(f"  {line}")
    print("  --- END PREVIEW ---")
    print(f"\n  Open the full report: {os.path.abspath(report_path)}")
    print(f"  Open the Word report: {os.path.abspath(docx_path)}")


if __name__ == "__main__":
    main()
